"""
Document ingestion — PDF, DOCX and scanned pages into page-addressable text.

Text is stored one row per page (DocumentPage) rather than as one blob, because
every downstream claim has to cite a page. An extracted figure whose provenance
is "somewhere in a 180-page sustainability report" is not evidence.

Pages that yield no text are recorded as image-only and routed to OCR. If OCR
is unavailable the page is marked NeedsOcr rather than empty, so a scanned
disclosure never silently reads as an absent one.
"""

import hashlib
import os
import uuid

from sqlalchemy import select

from esg import clock
from esg.config import settings
from esg.db.models import DocumentPage, EsgDocumentRegister
from esg.db.scope import require_principal
from esg.documents import ocr
from esg.security import audit, rbac

# Below this, a page is treated as having no real text layer.
MIN_TEXT_CHARS = 20

STATUS_PROCESSED = "Processed"
STATUS_PARTIAL = "ProcessedWithGaps"
STATUS_NEEDS_OCR = "NeedsOcr"
STATUS_FAILED = "Failed"


class IngestError(RuntimeError):
    pass


class UnsupportedDocument(IngestError):
    pass


def sha256_file(path):
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


# ── format readers ──

def _read_pdf_pages(path):
    """Yield (page_number, text, is_image_only, page_image_bytes_or_None).

    pdfplumber gives the text layer; PyMuPDF renders a raster for pages that
    have none, so OCR has something to work on.
    """
    import pdfplumber

    rendered = None
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if len(text) >= MIN_TEXT_CHARS:
                yield index, text, False, None
                continue

            if rendered is None:
                rendered = _open_with_pymupdf(path)
            image_bytes = _render_page(rendered, index - 1)
            yield index, text, True, image_bytes

    if rendered is not None:
        rendered.close()


def _open_with_pymupdf(path):
    import fitz

    return fitz.open(path)


def _render_page(document, page_index, zoom=2.0):
    import fitz

    try:
        page = document.load_page(page_index)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        return pixmap.tobytes("png")
    except Exception:
        return None


def _read_docx_pages(path):
    """DOCX has no page model, so paragraphs and tables are grouped into
    blocks of roughly a page. Block numbers are cited as page numbers and the
    document register records that they are logical, not printed, pages."""
    import docx

    document = docx.Document(path)
    chunks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    block, size, number = [], 0, 1
    for chunk in chunks:
        block.append(chunk)
        size += len(chunk)
        if size >= 2500:
            yield number, "\n".join(block), False, None
            block, size, number = [], 0, number + 1
    if block:
        yield number, "\n".join(block), False, None


def _read_text_pages(path):
    with open(path, "r", encoding="utf-8", errors="replace") as handle:
        content = handle.read()
    step = 3000
    if not content.strip():
        yield 1, "", True, None
        return
    for number, start in enumerate(range(0, len(content), step), start=1):
        yield number, content[start:start + step], False, None


_READERS = {
    ".pdf": _read_pdf_pages,
    ".docx": _read_docx_pages,
    ".txt": _read_text_pages,
    ".md": _read_text_pages,
}


def reader_for(path):
    extension = os.path.splitext(path)[1].lower()
    reader = _READERS.get(extension)
    if reader is None:
        raise UnsupportedDocument(
            f"{extension or path!r} is not a supported document type. "
            f"Supported: {', '.join(sorted(_READERS))}"
        )
    return reader


# ── ingestion ──

def ingest_document(session, path, deal_id, company_id, document_type=None,
                    reporting_year=None, document_name=None, confidentiality="Confidential",
                    source_system=None, attempt_ocr=True):
    """Register a document and store its text one page at a time.

    Re-ingesting identical content (same sha256) for the same deal returns the
    existing record rather than duplicating it — data rooms routinely serve the
    same file under several names.
    """
    principal = require_principal()
    rbac.check(rbac.INGEST_DATA, deal_id=deal_id, principal=principal)

    if not os.path.exists(path):
        raise IngestError(f"File not found: {path}")
    reader = reader_for(path)
    content_hash = sha256_file(path)

    existing = session.execute(
        select(EsgDocumentRegister).where(
            EsgDocumentRegister.content_sha256 == content_hash,
            EsgDocumentRegister.deal_id == deal_id,
        )
    ).scalars().first()
    if existing is not None:
        return existing

    retention_days = settings().retention_days_documents
    document = EsgDocumentRegister(
        document_id=uuid.uuid4().hex[:32],
        deal_id=deal_id,
        company_id=company_id,
        document_name=document_name or os.path.basename(path),
        document_type=document_type,
        reporting_year=reporting_year,
        source_system=source_system,
        file_path=os.path.abspath(path),
        content_sha256=content_hash,
        confidentiality_flag=confidentiality,
        processing_status="Processing",
        ingested_at=clock.now(),
        ingested_by=principal.username,
        retention_expires_at=(clock.today() + _days(retention_days)),
    )
    session.add(document)
    session.flush()

    ocr_backend = None
    pages_needing_ocr = 0
    page_count = 0

    try:
        for number, text, image_only, image_bytes in reader(path):
            page_count = number
            method = "text_layer"
            confidence = None

            if image_only:
                text, method, confidence, needed = _ocr_page(
                    image_bytes, attempt_ocr, ocr_backend_holder := [ocr_backend]
                )
                ocr_backend = ocr_backend_holder[0]
                pages_needing_ocr += needed

            session.add(DocumentPage(
                page_uid=f"{document.document_id}:{number}",
                deal_id=deal_id,
                document_id=document.document_id,
                page_number=number,
                text=text or None,
                char_count=len(text or ""),
                extraction_method=method,
                ocr_confidence=confidence,
                is_image_only=image_only,
            ))
    except Exception as exc:
        document.processing_status = STATUS_FAILED
        document.processing_error = f"{type(exc).__name__}: {exc}"
        audit.record(session, principal.username, "document.ingest_failed",
                     entity_type="esg_document_register", entity_id=document.document_id,
                     deal_id=deal_id, detail=document.processing_error)
        raise

    document.page_count = page_count
    if pages_needing_ocr and pages_needing_ocr == page_count:
        document.processing_status = STATUS_NEEDS_OCR
    elif pages_needing_ocr:
        document.processing_status = STATUS_PARTIAL
    else:
        document.processing_status = STATUS_PROCESSED
    document.processing_error = (
        f"{pages_needing_ocr} of {page_count} pages have no text layer and OCR is "
        "unavailable; these pages need manual review."
        if pages_needing_ocr else None
    )

    audit.record(
        session, principal.username, "document.ingested",
        entity_type="esg_document_register", entity_id=document.document_id,
        deal_id=deal_id,
        detail={"name": document.document_name, "pages": page_count,
                "status": document.processing_status,
                "pages_needing_ocr": pages_needing_ocr},
    )
    session.flush()
    return document


def _ocr_page(image_bytes, attempt_ocr, backend_holder):
    """Returns (text, method, confidence, needed_ocr_but_unavailable)."""
    if not attempt_ocr or image_bytes is None:
        return "", "image_only_skipped", None, 1

    backend = backend_holder[0]
    if backend is None:
        try:
            backend = ocr.get_backend()
        except ocr.OcrUnavailable:
            return "", "ocr_unavailable", None, 1
        backend_holder[0] = backend

    try:
        text, confidence = backend.image_to_text(image_bytes)
    except ocr.OcrUnavailable:
        return "", "ocr_unavailable", None, 1
    except Exception:
        return "", "ocr_failed", None, 1
    return text, f"ocr:{backend.name}", confidence, 0


def _days(count):
    from datetime import timedelta

    return timedelta(days=count)


def page_text(session, document_id, page_number):
    """Fetch one page — used to render a citation in the UI."""
    return session.execute(
        select(DocumentPage).where(
            DocumentPage.document_id == document_id,
            DocumentPage.page_number == page_number,
        )
    ).scalars().first()


def pages_needing_review(session, deal_id=None):
    """Pages with no usable text, so gaps are visible rather than assumed empty."""
    stmt = select(DocumentPage).where(DocumentPage.is_image_only.is_(True))
    if deal_id:
        stmt = stmt.where(DocumentPage.deal_id == deal_id)
    return session.execute(stmt.order_by(DocumentPage.page_uid)).scalars().all()
